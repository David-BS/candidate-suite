"""
Fills a cover-letter template with the provided data and inserts the signature.

Usage:
    python fill_cover_letter.py \\
        --language en|fr \\
        --template-path /path/to/template.docx \\
        --signature-path /path/to/signature_b64.txt \\
        --output-path /path/to/output.docx \\
        --data-json '<json_string>'

The --data-json JSON must contain the following keys:
    - sender_name, sender_street, sender_postal_code, sender_city
    - sender_email, sender_phone, sender_linkedin
    - sender_full_name
    - recruiter_name, recruiter_title (use "Hiring Manager" / "Madame, Monsieur" if unknown)
    - company_name
    - date_letter (format: "27 May 2026" or "27 mai 2026")
    - job_title
    - greeting (salutation, e.g.: "Dear Jane Smith," or "Madame, Monsieur,")
    - paragraph_1_intro, paragraph_2_current, paragraph_3_experience,
      paragraph_4_value, paragraph_5_closing

The script:
    1. Loads the template
    2. Replaces the placeholders (PRESERVING multiple runs)
    3. Inserts the signature image under "Sincerely,"/"Cordialement," and under the name
    4. Saves the .docx
"""

import argparse
import re


# Language-neutral sentinel for a required-but-unknown datum. The orchestrator
# passes it (instead of inventing a placeholder) to force a clean exit-2 when a
# mandatory field can't be filled; the script treats it as blank. This replaces
# the former multilingual placeholder word-list (FR/EN only — see LNG-2 S3b).
MISSING_SENTINEL = "__MISSING__"


def iso639_1(value):
    """Validate --language as an ISO 639-1 *form*: two lowercase letters.

    Open by design (LNG-1 L5) — no maintained list, no external dependency. Post-L6 there is no per-language lock: the .md generators validate the
    model-provided label *key set* (`--labels-json`), not the language, and the letter
    fills a single neutral template. `--language` is form-validated metadata only; any
    ISO code works.
    """
    if not re.fullmatch(r"[a-z]{2}", value or ""):
        raise argparse.ArgumentTypeError(
            f"--language must be a 2-letter ISO 639-1 code (lowercase), e.g. en, fr; got: {value!r}"
        )
    return value


import base64
import io
import json
import sys
from pathlib import Path

from docx import Document
from docx.shared import Cm


def decode_base64_signature(base64_string):
    """Decodes a base64 signature and returns a BytesIO."""
    base64_clean = base64_string.replace("\n", "").replace(" ", "").replace("\r", "")
    image_bytes = base64.b64decode(base64_clean)
    return io.BytesIO(image_bytes)


def replace_text_in_paragraph(paragraph, replacements):
    """Replaces the placeholders in a paragraph while PRESERVING multiple runs.

    This is crucial to keep differentiated formatting (e.g.: 2 bold lines + 1 normal line
    in the recipient block).
    """
    for key, value in replacements.items():
        # Check in each run individually
        for run in paragraph.runs:
            if key in run.text:
                # Replace only in THIS run (preserves the others' formatting)
                run.text = run.text.replace(key, value)


def _make_floating(run, cx, cy):
    """Turns a run's INLINE image into a FLOATING image (wp:anchor).

    Setting validated in Word: square wrap on both sides, overlap allowed,
    H position 143.75 pt from the column, V -0.3 pt from the anchor paragraph,
    wrap margins 9 pt left/right. The image thus leaves the text flow: no more
    unsightly gap with « Cordialement, » nor inflated line height.
    """
    from docx.oxml import parse_xml
    from docx.oxml.ns import qn, nsdecls

    PT = 12700  # EMU per point
    drawing = run._element.find(qn("w:drawing"))
    inline = drawing.find(qn("wp:inline"))
    graphic = inline.find(qn("a:graphic"))

    h_off = int(round(143.75 * PT))
    v_off = int(round(-0.3 * PT))
    dist = int(round(9 * PT))

    anchor = parse_xml(
        '<wp:anchor %s distT="0" distB="0" distL="%d" distR="%d" '
        'simplePos="0" relativeHeight="251658240" behindDoc="0" '
        'locked="0" layoutInCell="1" allowOverlap="1">'
        '<wp:simplePos x="0" y="0"/>'
        '<wp:positionH relativeFrom="column"><wp:posOffset>%d</wp:posOffset></wp:positionH>'
        '<wp:positionV relativeFrom="paragraph"><wp:posOffset>%d</wp:posOffset></wp:positionV>'
        '<wp:extent cx="%d" cy="%d"/>'
        '<wp:effectExtent l="0" t="0" r="0" b="0"/>'
        '<wp:wrapSquare wrapText="bothSides"/>'
        '<wp:docPr id="100" name="Signature"/>'
        "<wp:cNvGraphicFramePr/>"
        "</wp:anchor>"
        % (nsdecls("wp", "a", "r", "pic"), dist, dist, h_off, v_off, int(cx), int(cy))
    )
    anchor.append(graphic)  # move the graphic into the anchor
    drawing.replace(inline, anchor)  # inline -> floating


def insert_signature(doc, signature_base64):
    """Replaces the {{SIGNATURE_IMAGE}} placeholder with a FLOATING signature image
    (see _make_floating), anchored to the signatory-name paragraph.

    The floating image takes up no text line: the « Cordialement, / Name » block
    stays compact, and the image sits to the right without creating a gap.
    Size 2.6 x 1.8 cm. The placeholder paragraph is removed.
    """
    sig_bytes = decode_base64_signature(signature_base64)
    cx, cy = Cm(2.6), Cm(1.8)
    paras = doc.paragraphs

    for idx, paragraph in enumerate(paras):
        if "{{SIGNATURE_IMAGE}}" in paragraph.text:
            # Anchor paragraph = name (previous non-empty), otherwise the placeholder
            anchor_para = None
            for j in range(idx - 1, -1, -1):
                if paras[j].text.strip():
                    anchor_para = paras[j]
                    break
            host = anchor_para if anchor_para is not None else paragraph

            run = host.add_run()
            run.add_picture(sig_bytes, width=cx, height=cy)

            try:
                _make_floating(run, cx, cy)
            except Exception as e:
                # Fallback: if the floating conversion fails, the image stays inline
                print(
                    f"⚠️ Signature flottante impossible, repli inline : {e}",
                    file=sys.stderr,
                )

            if host is paragraph:
                # placeholder = host: remove the placeholder text, keep the image
                for r in list(paragraph.runs):
                    if "{{SIGNATURE_IMAGE}}" in r.text:
                        r.text = r.text.replace("{{SIGNATURE_IMAGE}}", "")
            else:
                # remove the now-useless placeholder paragraph
                el = paragraph._element
                el.getparent().remove(el)
            return True

    return False


def remove_signature_placeholder(doc):
    """Cleanly removes the {{SIGNATURE_IMAGE}} placeholder (letter without a signature image)."""
    for paragraph in doc.paragraphs:
        if "{{SIGNATURE_IMAGE}}" in paragraph.text:
            for run in paragraph.runs:
                if "{{SIGNATURE_IMAGE}}" in run.text:
                    run.text = run.text.replace("{{SIGNATURE_IMAGE}}", "")
            return True
    return False


def build_replacements(data):
    """Builds the {{PLACEHOLDER}} -> value replacement dict
    from the provided data dict.
    """
    # Mapping placeholder -> key in data
    mapping = {
        "{{SENDER_NAME}}": "sender_name",
        "{{SENDER_STREET}}": "sender_street",
        "{{SENDER_POSTAL_CODE}}": "sender_postal_code",
        "{{SENDER_CITY}}": "sender_city",
        "{{SENDER_EMAIL}}": "sender_email",
        "{{SENDER_LINKEDIN}}": "sender_linkedin",
        "{{SENDER_PHONE}}": "sender_phone",
        "{{SENDER_FULL_NAME}}": "sender_full_name",
        "{{RECRUITER_NAME}}": "recruiter_name",
        "{{RECRUITER_TITLE}}": "recruiter_title",
        "{{COMPANY_NAME}}": "company_name",
        "{{DATE_LETTER}}": "date_letter",
        "{{JOB_TITLE}}": "job_title",
        "{{GREETING}}": "greeting",
        "{{SUBJECT_LABEL}}": "subject_label",
        "{{CLOSING}}": "closing",
        "{{PARAGRAPH_1_INTRO}}": "paragraph_1_intro",
        "{{PARAGRAPH_2_CURRENT}}": "paragraph_2_current",
        "{{PARAGRAPH_3_EXPERIENCE}}": "paragraph_3_experience",
        "{{PARAGRAPH_4_ACHIEVEMENTS}}": "paragraph_4_value",  # legacy name in the template
        "{{PARAGRAPH_5_CLOSING}}": "paragraph_5_closing",
    }

    replacements = {}
    for placeholder, data_key in mapping.items():
        value = data.get(data_key, "")
        if value is None:
            value = ""
        replacements[placeholder] = str(value)

    return replacements


def fill_template(
    language, template_path, signature_path, signature_base64_arg, output_path, data
):
    """Fills the template with the data and inserts the signature (if provided).

    The signature can come from two (mutually exclusive) sources:
    - signature_base64_arg : base64 string passed directly (from memory).
    - signature_path : file containing the base64 signature (legacy / session upload).

    If neither is provided or the value is invalid, the {{SIGNATURE_IMAGE}}
    placeholder is cleanly removed (letter with no image signature).
    """
    # Check that the template exists
    if not Path(template_path).exists():
        print(f"❌ Template not found: {template_path}", file=sys.stderr)
        sys.exit(1)

    # Signature resolution: base64 argument takes priority, otherwise file
    signature_base64 = None
    if signature_base64_arg and signature_base64_arg.strip():
        signature_base64 = signature_base64_arg.strip()
    elif signature_path and Path(signature_path).exists():
        try:
            with open(signature_path, "r") as f:
                content = f.read().strip()
            if content:
                signature_base64 = content
        except Exception as e:
            print(f"⚠️ Signature unreadable, ignored: {e}", file=sys.stderr)

    # #1 — recruiter name: the model provides it (a real name, or a localized generic
    #      such as "Service Recrutement" / "Recruitment Department" / "Personalabteilung"
    #      when unknown). No hardcoded language default here (agnostic — LNG-1/L6).

    # Load the template
    doc = Document(template_path)

    # Build the replacements
    replacements = build_replacements(data)

    # 1. Replace all placeholders EXCEPT the signature (preserving the runs)
    for paragraph in doc.paragraphs:
        if "{{SIGNATURE_IMAGE}}" not in paragraph.text:
            replace_text_in_paragraph(paragraph, replacements)

    # Also replace in tables if present
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if "{{SIGNATURE_IMAGE}}" not in paragraph.text:
                        replace_text_in_paragraph(paragraph, replacements)

    # #3bis — unknown recruiter (empty title): clean the orphan comma in the
    # recipient block (« Service Recrutement, » → « Service Recrutement »), without
    # touching line breaks (separate <w:br>) or the date line.
    if not str(data.get("recruiter_title") or "").strip():
        from docx.oxml.ns import qn as _qn

        rec_name = str(data.get("recruiter_name") or "")
        comp = str(data.get("company_name") or "")
        for p in doc.paragraphs:
            if rec_name and comp and rec_name in p.text and comp in p.text:
                for r in p.runs:
                    if r._element.find(_qn("w:br")) is not None:
                        t = r._element.find(_qn("w:t"))
                        if t is not None and t.text and t.text.rstrip().endswith(","):
                            t.text = t.text.rstrip()[:-1].rstrip()
                break

    # 2. Insert the signature if available and valid; otherwise remove the placeholder
    if signature_base64:
        try:
            inserted = insert_signature(doc, signature_base64)
            if not inserted:
                print(
                    "⚠️ Placeholder {{SIGNATURE_IMAGE}} not found in the template",
                    file=sys.stderr,
                )
        except Exception as e:
            print(f"⚠️ Invalid signature, placeholder removed: {e}", file=sys.stderr)
            remove_signature_placeholder(doc)
    else:
        # No signature: cleanly remove the placeholder so it isn't displayed
        remove_signature_placeholder(doc)
        print(
            "ℹ️ No signature provided: letter generated without a signature image.",
            file=sys.stderr,
        )

    # 3. Save
    Path(output_path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)

    print(f"✅ Cover letter generated: {output_path}")
    return output_path


# ── 0.4.0 — "one page" guardrail (LIV-1) ─────────────────────────────────────
# HARD cap on body length (sum of the 5 paragraphs, spaces included).
# Calibrated and validated in Word: 2800 characters fit on one A4 page for
# FR AND EN (structurally identical templates; EN is ~13% more compact).
# Beyond that, the letter overflows to a 2nd page. Adjustable via --body-cap.
BODY_CAP_DEFAULT = 2800

# Body paragraphs (order) and SOFT distribution ratios (15/22/26/22/15).
# Individually non-blocking: they indicate WHICH paragraph to shorten.
BODY_PARAGRAPHS = [
    ("paragraph_1_intro", 0.15),
    ("paragraph_2_current", 0.22),
    ("paragraph_3_experience", 0.26),
    ("paragraph_4_value", 0.22),
    ("paragraph_5_closing", 0.15),
]


def check_body_cap(data, cap):
    """Measures the body and compares it to the cap.
    Returns (ok, total, details); details = [(field, length, target, over_target)]."""
    details = []
    total = 0
    for field, ratio in BODY_PARAGRAPHS:
        length = len(str(data.get(field, "") or ""))
        total += length
        target = int(round(cap * ratio))
        details.append((field, length, target, length > target * 1.2))
    return total <= cap, total, details


def main():
    parser = argparse.ArgumentParser(description="Fills a cover-letter template")
    parser.add_argument(
        "--language",
        required=True,
        type=iso639_1,
        metavar="LANG",
        help="ISO 639-1 language code (lowercase), e.g. en, fr",
    )
    parser.add_argument(
        "--template-path", required=True, help="Path to the .docx template"
    )
    parser.add_argument(
        "--signature-path",
        default="",
        help="Path to the base64 signature file (optional)",
    )
    parser.add_argument(
        "--signature-base64",
        default="",
        help="Base64 signature (direct string, optional; takes priority over --signature-path)",
    )
    parser.add_argument("--output-path", required=True, help="Output .docx path")
    parser.add_argument("--data-json", required=True, help="JSON with all the data")
    parser.add_argument(
        "--body-cap",
        type=int,
        default=BODY_CAP_DEFAULT,
        help=f"Body cap in characters (default {BODY_CAP_DEFAULT})",
    )

    args = parser.parse_args()

    # Parse the JSON
    try:
        data = json.loads(args.data_json)
    except json.JSONDecodeError as e:
        print(f"❌ Invalid JSON: {e}", file=sys.stderr)
        sys.exit(1)

    # Check the required fields
    required_fields = [
        "sender_name",
        "sender_street",
        "sender_postal_code",
        "sender_city",
        "sender_email",
        "sender_phone",
        "sender_linkedin",
        "sender_full_name",
        "recruiter_name",
        "recruiter_title",
        "company_name",
        "date_letter",
        "job_title",
        "greeting",
        "paragraph_1_intro",
        "paragraph_2_current",
        "paragraph_3_experience",
        "paragraph_4_value",
        "paragraph_5_closing",
    ]

    # Detect missing fields (absent / None) and blank values. A value is "blank"
    # if it is empty/whitespace OR equals the MISSING sentinel — a language-neutral
    # token the orchestrator passes when a required datum is unknown. This is a
    # STRUCTURAL contract (LNG-2 S3b): it replaces the former multilingual
    # placeholder word-list ("à compléter"/"to complete"/… — FR/EN only, brittle
    # in other languages). The rule remains: never invent a placeholder; if a
    # mandatory datum is missing, ASK the user (see the GUIDE's "Mandatory data").
    def _is_blank(v):
        if v is None:
            return True
        s = str(v).strip()
        return (not s) or (s == MISSING_SENTINEL)

    missing = [f for f in required_fields if f not in data or data[f] is None]
    blank = [
        f
        for f in required_fields
        if f in data and data[f] is not None and _is_blank(data[f])
    ]

    # Critical fields for which a blank value (empty or sentinel) is unacceptable.
    critical_fields = [
        "sender_full_name",
        "sender_street",
        "sender_postal_code",
        "sender_city",
        "sender_email",
        "company_name",
        "paragraph_1_intro",
        "paragraph_2_current",
        "paragraph_3_experience",
        "paragraph_4_value",
        "paragraph_5_closing",
    ]
    critical_blank = [f for f in critical_fields if f in blank]

    if missing:
        print(f"\u274c Missing fields in the JSON: {missing}", file=sys.stderr)
        sys.exit(1)
    if critical_blank:
        print(
            "\u274c Mandatory data missing (empty or the '__MISSING__' sentinel) \u2014 refusing to generate an incomplete letter:",
            file=sys.stderr,
        )
        for f in critical_blank:
            print(f"   - {f}", file=sys.stderr)
        print(
            "   \u2192 Do NOT fill these with a placeholder. Ask the user for them, then re-run",
            file=sys.stderr,
        )
        print('     (see the GUIDE\'s "Mandatory data" section).', file=sys.stderr)
        sys.exit(2)
    non_critical_blank = [f for f in blank if f not in critical_fields]
    if non_critical_blank:
        # Blank but non-critical fields (e.g. linkedin): warn without blocking.
        print(
            f"\u26a0\ufe0f Blank fields (non-blocking): {non_critical_blank}",
            file=sys.stderr,
        )

    # 0.4.0 — "one page" guardrail (LIV-1): reject if the body exceeds the cap.
    ok, total, details = check_body_cap(data, args.body_cap)
    if not ok:
        # Paragraph to shorten FIRST = the one most over its target (ratio).
        prio = None
        worst = 1.0
        for field, length, target, _ in details:
            if target and length > target and length / target > worst:
                worst = length / target
                prio = field
        print(
            f"❌ Letter too long: body = {total} characters (cap {args.body_cap}). "
            f"It would overflow onto a 2nd page.",
            file=sys.stderr,
        )
        print(
            "   Shorten the CONTENT (never the margins or font), then regenerate.",
            file=sys.stderr,
        )
        print("   Per-paragraph sizes — target = ratio × cap:", file=sys.stderr)
        for field, length, target, _ in details:
            mark = "  ⟵ shorten first" if field == prio else ""
            print(f"     - {field}: {length} (target ~{target}){mark}", file=sys.stderr)
        sys.exit(2)
    over = [d for d in details if d[3]]
    if over:
        print(
            "⚠️ Balancing (non-blocking) — paragraph(s) well above their target:",
            file=sys.stderr,
        )
        for field, length, target, _ in over:
            print(f"     - {field}: {length} (target ~{target})", file=sys.stderr)

    fill_template(
        language=args.language,
        template_path=args.template_path,
        signature_path=args.signature_path,
        signature_base64_arg=args.signature_base64,
        output_path=args.output_path,
        data=data,
    )


if __name__ == "__main__":
    main()
