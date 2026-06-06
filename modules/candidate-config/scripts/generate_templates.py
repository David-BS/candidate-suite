"""
Script that generates the Cover Letter templates (EN + FR).

Usage:
    python generate_templates.py [--style hybrid|block] [output_dir]

Arguments:
    --style : hybrid (default) or block
    output_dir : output folder (default /mnt/user-data/outputs/)

Header styles:
- hybrid : 3 compact lines (name, full address, digital coordinates with |)
- block  : 6 classic lines (name, street, postal code+city, email, LinkedIn, phone)

Common characteristics:
- Gray border (Single, light gray, 1pt) to the left of the address block
- Uniform Calibri (Sans Serif) fonts
- Recipient block: 2 bold lines, 3rd line (city + date) normal
- Handwritten signature (placeholder) under the signatory's name
- Sincerely/Cordialement in bold
"""

import argparse
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import parse_xml


def add_left_border(paragraph):
    """Adds a grey left border (Single, 1pt, light grey) with spacing"""
    pPr = paragraph._element.get_or_add_pPr()
    existing_pBdr = pPr.find(qn("w:pBdr"))
    if existing_pBdr is not None:
        pPr.remove(existing_pBdr)
    pBdr_xml = """<w:pBdr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
        <w:left w:val="single" w:sz="8" w:space="10" w:color="C0C0C0"/>
    </w:pBdr>"""
    pBdr = parse_xml(pBdr_xml)
    pPr.append(pBdr)


def add_sender_block_hybrid(doc):
    """Adds the sender address block in HYBRID mode (3 lines)"""
    # Line 1: Name (14pt, bold, blue)
    p_sender = doc.add_paragraph()
    add_left_border(p_sender)
    p_sender.paragraph_format.space_after = Pt(0)
    p_sender.paragraph_format.line_spacing = 1.0
    run = p_sender.add_run("{{SENDER_NAME}}")
    run.font.name = "Calibri"
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = RGBColor(31, 73, 125)

    # Line 2: full address (street, postal code, city)
    p_addr = doc.add_paragraph()
    add_left_border(p_addr)
    p_addr.paragraph_format.space_after = Pt(0)
    p_addr.paragraph_format.line_spacing = 1.0
    run_addr = p_addr.add_run(
        "{{SENDER_STREET}}, {{SENDER_POSTAL_CODE}} {{SENDER_CITY}}"
    )
    run_addr.font.name = "Calibri"
    run_addr.font.size = Pt(10)
    run_addr.font.bold = False

    # Line 3: digital coordinates (email | LinkedIn | phone)
    p_digital = doc.add_paragraph()
    add_left_border(p_digital)
    p_digital.paragraph_format.space_after = Pt(0)
    p_digital.paragraph_format.line_spacing = 1.0

    # Email (blue underlined)
    run_email = p_digital.add_run("{{SENDER_EMAIL}}")
    run_email.font.name = "Calibri"
    run_email.font.size = Pt(10)
    run_email.font.color.rgb = RGBColor(0, 0, 255)
    run_email.font.underline = True
    run_email.font.bold = False

    # Separator
    run_sep1 = p_digital.add_run(" | ")
    run_sep1.font.name = "Calibri"
    run_sep1.font.size = Pt(10)
    run_sep1.font.bold = False

    # LinkedIn (blue underlined)
    run_linkedin = p_digital.add_run("{{SENDER_LINKEDIN}}")
    run_linkedin.font.name = "Calibri"
    run_linkedin.font.size = Pt(10)
    run_linkedin.font.color.rgb = RGBColor(0, 0, 255)
    run_linkedin.font.underline = True
    run_linkedin.font.bold = False

    # Separator
    run_sep2 = p_digital.add_run(" | ")
    run_sep2.font.name = "Calibri"
    run_sep2.font.size = Pt(10)
    run_sep2.font.bold = False

    # Phone (normal color)
    run_phone = p_digital.add_run("{{SENDER_PHONE}}")
    run_phone.font.name = "Calibri"
    run_phone.font.size = Pt(10)
    run_phone.font.bold = False


def add_sender_block_block(doc):
    """Adds the sender address block in BLOCK mode (6 lines)"""
    # Line 1: Name (14pt, bold, blue)
    p_sender = doc.add_paragraph()
    add_left_border(p_sender)
    p_sender.paragraph_format.space_after = Pt(0)
    p_sender.paragraph_format.line_spacing = 1.0
    run = p_sender.add_run("{{SENDER_NAME}}")
    run.font.name = "Calibri"
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = RGBColor(31, 73, 125)

    # Line 2: Street
    p_addr1 = doc.add_paragraph("{{SENDER_STREET}}")
    add_left_border(p_addr1)
    p_addr1.paragraph_format.space_after = Pt(0)
    p_addr1.paragraph_format.line_spacing = 1.0
    p_addr1.runs[0].font.name = "Calibri"
    p_addr1.runs[0].font.size = Pt(10)
    p_addr1.runs[0].font.bold = False

    # Line 3: Postal code + city
    p_addr2 = doc.add_paragraph("{{SENDER_POSTAL_CODE}} {{SENDER_CITY}}")
    add_left_border(p_addr2)
    p_addr2.paragraph_format.space_after = Pt(0)
    p_addr2.paragraph_format.line_spacing = 1.0
    p_addr2.runs[0].font.name = "Calibri"
    p_addr2.runs[0].font.size = Pt(10)
    p_addr2.runs[0].font.bold = False

    # Line 4: Email (blue underlined)
    p_email = doc.add_paragraph()
    add_left_border(p_email)
    p_email.paragraph_format.space_after = Pt(0)
    p_email.paragraph_format.line_spacing = 1.0
    run_email = p_email.add_run("{{SENDER_EMAIL}}")
    run_email.font.name = "Calibri"
    run_email.font.size = Pt(10)
    run_email.font.color.rgb = RGBColor(0, 0, 255)
    run_email.font.underline = True
    run_email.font.bold = False

    # Line 5: LinkedIn (blue underlined)
    p_linkedin = doc.add_paragraph()
    add_left_border(p_linkedin)
    p_linkedin.paragraph_format.space_after = Pt(0)
    p_linkedin.paragraph_format.line_spacing = 1.0
    run_linkedin = p_linkedin.add_run("{{SENDER_LINKEDIN}}")
    run_linkedin.font.name = "Calibri"
    run_linkedin.font.size = Pt(10)
    run_linkedin.font.color.rgb = RGBColor(0, 0, 255)
    run_linkedin.font.underline = True
    run_linkedin.font.bold = False

    # Line 6: Phone
    p_phone = doc.add_paragraph("{{SENDER_PHONE}}")
    add_left_border(p_phone)
    p_phone.paragraph_format.space_after = Pt(0)
    p_phone.paragraph_format.line_spacing = 1.0
    p_phone.runs[0].font.name = "Calibri"
    p_phone.runs[0].font.size = Pt(10)
    p_phone.runs[0].font.bold = False


def add_common_body(doc):
    """Adds the shared body (recipient, subject, paragraphs, signature)"""
    # Blank line
    doc.add_paragraph()

    # Recipient block + date
    p_recipient = doc.add_paragraph()
    p_recipient.paragraph_format.left_indent = Inches(3.4)

    run_rec = p_recipient.add_run("{{RECRUITER_NAME}}, {{RECRUITER_TITLE}}\n")
    run_rec.font.name = "Calibri"
    run_rec.font.size = Pt(11)
    run_rec.font.bold = True

    run_company = p_recipient.add_run("{{COMPANY_NAME}}\n\n")
    run_company.font.name = "Calibri"
    run_company.font.size = Pt(11)
    run_company.font.bold = True

    run_location = p_recipient.add_run("{{SENDER_CITY}}, {{DATE_LETTER}}")
    run_location.font.name = "Calibri"
    run_location.font.size = Pt(11)
    run_location.font.bold = False

    # Subject line
    p_subject = doc.add_paragraph()
    p_subject.paragraph_format.space_before = Pt(12)
    p_subject.paragraph_format.space_after = Pt(0)
    # Neutral subject: the model supplies {{SUBJECT_LABEL}} = localized label + separator
    # (e.g. 'Poste\u00a0: ' / 'Position: ' / 'Betreff: ') in the target language.
    subject_text = "{{SUBJECT_LABEL}}{{JOB_TITLE}}"
    run_subject = p_subject.add_run(subject_text)
    run_subject.font.name = "Calibri"
    run_subject.font.size = Pt(11)
    run_subject.font.bold = True

    # Salutation
    greeting_placeholder = "{{GREETING}}"
    p_greeting = doc.add_paragraph(greeting_placeholder)
    p_greeting.paragraph_format.space_before = Pt(6)
    p_greeting.paragraph_format.space_after = Pt(0)
    p_greeting.runs[0].font.name = "Calibri"
    p_greeting.runs[0].font.size = Pt(11)

    # Body: 5 justified paragraphs
    para_texts = [
        "{{PARAGRAPH_1_INTRO}}",
        "{{PARAGRAPH_2_CURRENT}}",
        "{{PARAGRAPH_3_EXPERIENCE}}",
        "{{PARAGRAPH_4_ACHIEVEMENTS}}",
        "{{PARAGRAPH_5_CLOSING}}",
    ]
    for para_text in para_texts:
        p = doc.add_paragraph(para_text)
        p.paragraph_format.space_after = Pt(3.6)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.runs[0].font.name = "Calibri"
        p.runs[0].font.size = Pt(11)

    # Closing
    closing_text = "{{CLOSING}}"
    p_closing = doc.add_paragraph(closing_text)
    p_closing.paragraph_format.space_before = Pt(12)
    p_closing.runs[0].font.name = "Calibri"
    p_closing.runs[0].font.size = Pt(11)
    p_closing.runs[0].font.bold = True

    # Full name
    p_name = doc.add_paragraph("{{SENDER_FULL_NAME}}")
    p_name.paragraph_format.space_after = Pt(0)
    p_name.paragraph_format.space_before = Pt(6)
    p_name.runs[0].font.name = "Calibri"
    p_name.runs[0].font.bold = True
    p_name.runs[0].font.size = Pt(11)

    # Handwritten signature (placeholder)
    p_sig = doc.add_paragraph()
    p_sig.paragraph_format.space_after = Pt(0)
    p_sig.paragraph_format.space_before = Pt(3)
    p_sig.add_run("{{SIGNATURE_IMAGE}}")


def create_template(style="hybrid"):
    """Creates the neutral Cover Letter template

    Args:
        style : 'hybrid' or 'block'
    """
    doc = Document()

    # A4 margins
    section = doc.sections[0]
    section.top_margin = Cm(1.25)
    section.bottom_margin = Cm(1)
    section.left_margin = Cm(1.75)
    section.right_margin = Cm(1.5)

    # Address block depending on the style
    if style == "hybrid":
        add_sender_block_hybrid(doc)
    elif style == "block":
        add_sender_block_block(doc)
    else:
        raise ValueError(f"Style inconnu : {style} (choisir 'hybrid' ou 'block')")

    # Shared body
    add_common_body(doc)

    return doc


def main():
    parser = argparse.ArgumentParser(
        description="Generates the neutral Cover Letter template (language-agnostic)"
    )
    parser.add_argument(
        "--style",
        choices=["hybrid", "block"],
        default="hybrid",
        help="Header style: hybrid (3 lines) or block (6 lines)",
    )
    parser.add_argument(
        "output_dir", nargs="?", default="/mnt/user-data/outputs/", help="Output folder"
    )

    args = parser.parse_args()

    output_dir = Path(args.output_dir)
    output_dir.mkdir(parents=True, exist_ok=True)

    # Single neutral template (language-agnostic — L6). Subject label, greeting and
    # closing are placeholders filled by the model in the target language.
    doc = create_template(style=args.style)
    output = output_dir / "Cover_letter_template.docx"
    doc.save(str(output))
    print(f"✅ Neutral template created (style {args.style}): {output}")

    print(f"\n📋 Style: {args.style}")
    if args.style == "hybrid":
        print(
            "  - 3 lines: name / full address / digital contacts (email | linkedin | phone)"
        )
    else:
        print(
            "  - 6 lines: name / street / postcode+city / email / linkedin / phone"
        )

    print("\n📋 Supported placeholders:")
    placeholders = [
        "{{SENDER_NAME}}",
        "{{SENDER_STREET}}",
        "{{SENDER_POSTAL_CODE}}",
        "{{SENDER_CITY}}",
        "{{SENDER_EMAIL}}",
        "{{SENDER_LINKEDIN}}",
        "{{SENDER_PHONE}}",
        "{{RECRUITER_NAME}}",
        "{{RECRUITER_TITLE}}",
        "{{COMPANY_NAME}}",
        "{{DATE_LETTER}}",
        "{{JOB_TITLE}}",
        "{{SUBJECT_LABEL}}",
        "{{GREETING}}",
        "{{CLOSING}}",
        "{{PARAGRAPH_1_INTRO}}",
        "{{PARAGRAPH_2_CURRENT}}",
        "{{PARAGRAPH_3_EXPERIENCE}}",
        "{{PARAGRAPH_4_ACHIEVEMENTS}}",
        "{{PARAGRAPH_5_CLOSING}}",
        "{{SIGNATURE_IMAGE}}",
        "{{SENDER_FULL_NAME}}",
    ]
    for ph in placeholders:
        print(f"  • {ph}")


if __name__ == "__main__":
    main()
